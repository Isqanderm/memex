from src.adapters.docx import DOCX_MIME, DocxAdapter
from src.adapters.protocol import Source


def test_docx_can_handle_mime():
    assert DocxAdapter().can_handle(Source(path="doc.docx", mime_type=DOCX_MIME))

def test_docx_can_handle_by_extension():
    assert DocxAdapter().can_handle(Source(path="doc.docx", mime_type="application/octet-stream", filename="doc.docx"))

def test_docx_cannot_handle_pdf():
    assert not DocxAdapter().can_handle(Source(path="doc.pdf", mime_type="application/pdf"))

def test_docx_parses_paragraphs(tmp_path):
    from docx import Document as DocxDoc
    d = DocxDoc()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    path = tmp_path / "test.docx"
    d.save(str(path))

    doc = DocxAdapter().parse(Source(path=str(path), mime_type=DOCX_MIME))
    contents = " ".join(s.content for s in doc.sections)
    assert "First paragraph" in contents

def test_docx_extracts_headings(tmp_path):
    from docx import Document as DocxDoc
    d = DocxDoc()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("Intro text.")
    d.add_heading("Details", level=2)
    d.add_paragraph("Detail text.")
    path = tmp_path / "headed.docx"
    d.save(str(path))

    doc = DocxAdapter().parse(Source(path=str(path), mime_type=DOCX_MIME))
    headings = [s.heading for s in doc.sections if s.heading]
    assert len(headings) >= 1
