from src.adapters.protocol import Source
from src.models.parsed import ParsedDocument, Section

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
                  "Title": 1, "Subtitle": 2}


class DocxAdapter:
    def can_handle(self, source: Source) -> bool:
        path = source.filename or source.path
        return source.mime_type == DOCX_MIME or path.endswith(".docx")

    def parse(self, source: Source) -> ParsedDocument:
        from docx import Document as DocxDocument
        doc = DocxDocument(source.path)
        sections = []
        current_heading = None
        current_level = 0
        current_paragraphs: list[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            level = HEADING_STYLES.get(style_name)

            if level is not None:
                if current_paragraphs:
                    sections.append(Section(
                        content="\n".join(current_paragraphs),
                        heading=current_heading,
                        level=current_level,
                    ))
                    current_paragraphs = []
                current_heading = para.text.strip()
                current_level = level
            else:
                text = para.text.strip()
                if text:
                    current_paragraphs.append(text)

        if current_paragraphs:
            sections.append(Section(
                content="\n".join(current_paragraphs),
                heading=current_heading,
                level=current_level,
            ))

        return ParsedDocument(
            source=source.path,
            mime_type=DOCX_MIME,
            sections=sections or [Section(content="")],
            metadata={"filename": source.filename or source.path},
        )
